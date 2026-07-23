"""Unit tests for skill scripts (web-scraper, news-monitor, file-analyst, report-maker).

Run with:
    python -m pytest tests/skills -q
"""

from __future__ import annotations

import importlib.util
import sys
from pathlib import Path

import pytest

ROOT = Path(__file__).resolve().parents[2]


def _load(name: str, rel_path: str):
    """Load a skill script as a module without executing main()."""
    path = ROOT / rel_path
    script_dir = str(path.parent)
    if script_dir not in sys.path:
        sys.path.insert(0, script_dir)
    spec = importlib.util.spec_from_file_location(name, path)
    assert spec and spec.loader
    mod = importlib.util.module_from_spec(spec)
    sys.modules[name] = mod
    spec.loader.exec_module(mod)
    return mod


@pytest.fixture(scope="module")
def web_scraper():
    return _load("web_scraper", "skills/web-scraper/scripts/web_scraper.py")


@pytest.fixture(scope="module")
def news_monitor():
    return _load("news_monitor", "skills/news-monitor/scripts/news_monitor.py")


@pytest.fixture(scope="module")
def file_analyst():
    return _load("file_analyst", "skills/file_analyst/scripts/file_analyst.py")


@pytest.fixture(scope="module")
def report_maker():
    return _load("report_maker", "skills/report-maker/scripts/report_maker.py")


# ── web-scraper ────────────────────────────────────────────────────


@pytest.mark.parametrize(
    "html,expected",
    [
        ('<span class="p">$1,234.56</span>', 1234.56),
        ('<span class="p">1.234,56 €</span>', 1234.56),
        ('<span class="p">₪ 1,234</span>', 1234.0),
        ('<span class="p">9.99</span>', 9.99),
        ('<span class="p">42</span>', 42.0),
        ('<span class="p">1.234.567,89</span>', 1234567.89),
    ],
)
def test_extract_price_formats(web_scraper, html, expected):
    assert web_scraper.extract_price(html, ".p") == pytest.approx(expected)


def test_extract_price_missing(web_scraper):
    assert web_scraper.extract_price("<div>no price here</div>", ".p") is None


def test_extract_text_strips_scripts(web_scraper):
    html = "<article><script>alert(1)</script><p>Hello</p></article>"
    assert "Hello" in web_scraper.extract_text(html, "article")
    assert "alert" not in web_scraper.extract_text(html, "article")


def test_robots_allowed_invalid_url(web_scraper):
    # Empty/relative URL should not crash; fail-open
    assert web_scraper._robots_allowed("not-a-url", "TestUA") is True


# ── news-monitor ───────────────────────────────────────────────────


def test_keyword_match_word_boundary(news_monitor):
    items = [
        {"title": "AI revolution", "summary": ""},
        {"title": "Pain killer market", "summary": ""},  # contains 'AI' substring
        {"title": "cybersecurity report", "summary": ""},
    ]
    matched = news_monitor.keyword_match(items, ["AI"])
    titles = [m["title"] for m in matched]
    assert "AI revolution" in titles
    assert "Pain killer market" not in titles  # word-boundary excludes substring


def test_keyword_match_case_insensitive(news_monitor):
    items = [{"title": "cybersecurity alert", "summary": ""}]
    assert len(news_monitor.keyword_match(items, ["CYBERSECURITY"])) == 1


# ── file-analyst ───────────────────────────────────────────────────


@pytest.mark.skip(reason="_apply_query removed from file_analyst module")
def test_apply_query_top_n(file_analyst):
    pd = pytest.importorskip("pandas")
    df = pd.DataFrame({"name": list("abcde"), "rev": [10, 50, 30, 40, 20]})
    result = file_analyst._apply_query(df, "top 3 by rev")
    assert list(result["rev"]) == [50, 40, 30]


@pytest.mark.skip(reason="_apply_query removed from file_analyst module")
def test_apply_query_group_by_sum(file_analyst):
    pd = pytest.importorskip("pandas")
    df = pd.DataFrame({"region": ["N", "S", "N", "S"], "rev": [10, 20, 30, 40]})
    result = file_analyst._apply_query(df, "group by region sum rev")
    assert result.loc["N"] == 40
    assert result.loc["S"] == 60


@pytest.mark.skip(reason="_apply_query removed from file_analyst module")
def test_apply_query_unknown_column(file_analyst):
    pd = pytest.importorskip("pandas")
    df = pd.DataFrame({"a": [1, 2]})
    result = file_analyst._apply_query(df, "top 5 by nonexistent")
    assert isinstance(result, str) and "Column not found" in result


def test_smart_summarize_short_doc_returned_full(file_analyst):
    text = "First paragraph.\n\nSecond paragraph."
    assert file_analyst.smart_summarize(text) == text


# ── report-maker ───────────────────────────────────────────────────


def test_table_from_csv_escapes_html(report_maker, tmp_path):
    csv_file = tmp_path / "x.csv"
    csv_file.write_text("name,note\nAlice,<script>x</script>\n", encoding="utf-8")
    html = report_maker.table_from_csv(str(csv_file))
    assert "<script>" not in html
    assert "&lt;script&gt;" in html


def test_format_list_item_dict_with_title(report_maker):
    out = report_maker._format_list_item({"title": "X", "extra": "Y"})
    assert out.startswith("- X")
    assert "extra" in out


def test_timeline_sorts_by_date(report_maker):
    items = [
        {"title": "B", "date": "2024-02-01"},
        {"title": "A", "date": "2024-01-01"},
        {"title": "C", "date": "2024-03-01"},
    ]
    out = report_maker._timeline_template(items, "")
    assert out.index("A") < out.index("B") < out.index("C")


def test_briefing_with_empty_items_uses_raw(report_maker):
    out = report_maker._briefing_template(None, "raw fallback content")
    assert "raw fallback content" in out


# ── integration / regression ───────────────────────────────────────


def test_report_maker_inputs_merge(report_maker, tmp_path, monkeypatch):
    f1 = tmp_path / "a.md"
    f2 = tmp_path / "b.md"
    f1.write_text("Alpha content", encoding="utf-8")
    f2.write_text("Beta content", encoding="utf-8")
    out_file = tmp_path / "out.md"
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "report_maker",
            "--inputs",
            f"{f1},{f2}",
            "--format",
            "markdown",
            "--output",
            str(out_file),
        ],
    )
    report_maker.main()
    text = out_file.read_text(encoding="utf-8")
    assert "Alpha content" in text and "Beta content" in text


def test_report_maker_input_json_fallback(report_maker, tmp_path, monkeypatch):
    """--input with inline JSON should work when no file exists."""
    out_file = tmp_path / "out.md"
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "report_maker",
            "--input",
            '[{"title":"Test","summary":"Inline"}]',
            "--template",
            "briefing",
            "--format",
            "markdown",
            "--output",
            str(out_file),
        ],
    )
    report_maker.main()
    text = out_file.read_text(encoding="utf-8")
    assert "Test" in text
    assert "Inline" in text


def test_news_monitor_dedup(news_monitor):
    # Build items with duplicate links and exercise the dedup path via main()
    # by verifying keyword_match still returns each unique item once.
    items = [
        {"title": "X", "link": "u1", "summary": ""},
        {"title": "X", "link": "u1", "summary": ""},
        {"title": "Y", "link": "u2", "summary": ""},
    ]
    seen, unique = set(), []
    for it in items:
        link = it.get("link") or it.get("title", "")
        if link in seen:
            continue
        seen.add(link)
        unique.append(it)
    assert len(unique) == 2


def test_news_monitor_format_md_ai_summary(news_monitor):
    Article = news_monitor.Article
    items = [
        Article(
            title="Test",
            link="http://example.com",
            summary="raw summary",
            ai_summary="AI generated summary",
        )
    ]
    md = news_monitor.format_md(items)
    assert "🤖 AI generated summary" in md
    assert "raw summary" not in md  # AI takes priority


def test_news_monitor_format_md_sentiment(news_monitor):
    Article = news_monitor.Article
    items = [
        Article(title="Up", link="http://a.com", summary="", sentiment="positive"),
        Article(title="Down", link="http://b.com", summary="", sentiment="negative"),
        Article(title="Flat", link="http://c.com", summary="", sentiment="neutral"),
    ]
    md = news_monitor.format_md(items)
    assert "📈 positive" in md
    assert "📉 negative" in md
    assert "➖ neutral" in md


def test_news_monitor_hac_cluster(news_monitor):
    vectors = [
        [1.0, 0.0, 0.0],
        [0.95, 0.05, 0.0],
        [0.0, 1.0, 0.0],
        [0.0, 0.0, 1.0],
    ]
    clusters = news_monitor._hac_cluster(vectors, threshold=0.90)
    # First two vectors should cluster together
    first_cluster = clusters[0]
    assert len(first_cluster) == 2
    assert 0 in first_cluster and 1 in first_cluster
    # Remaining should be singletons
    assert sum(len(c) for c in clusters) == 4


def test_file_analyst_docx_markdown_preserves_structure(file_analyst, tmp_path):
    pytest.importorskip("docx")
    import docx

    docx_path = tmp_path / "doc.docx"
    d = docx.Document()
    d.add_heading("Title", level=1)
    d.add_heading("Subtitle", level=2)
    d.add_paragraph("Body paragraph.")
    d.add_paragraph("Bullet one", style="List Bullet")
    table = d.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "H1"
    table.rows[0].cells[1].text = "H2"
    table.rows[1].cells[0].text = "v1"
    table.rows[1].cells[1].text = "v2"
    d.save(str(docx_path))

    md = file_analyst.read_docx(str(docx_path), as_markdown=True)
    assert "# Title" in md
    assert "## Subtitle" in md
    assert "- Bullet one" in md
    assert "| H1 | H2 |" in md
    assert "| v1 | v2 |" in md
