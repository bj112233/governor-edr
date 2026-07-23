"""
File readers — PDF, DOCX, CSV, XLSX, JSON, TXT.

PDF extraction backends extracted to _pdf_readers.py (SRP).
"""
import json
import os
import sys
from pathlib import Path

# Ensure project root is on sys.path so we can import services.security_utils
# (this script runs in a subprocess with cwd=skill_dir, not project root)
# NOTE: skills must NOT import from services (subprocess isolation contract).
# _untrusted_wrap.py is a self-contained local copy of the sandboxing logic.

from _untrusted_wrap import wrap_untrusted_content

from _ocr_pdf import _MAX_PDF_BYTES, _MAX_PDF_PAGES
from _pdf_readers import (
    _check_pdf_size,
    _extract_pdfplumber,
    _extract_pymupdf,
    _extract_pypdf2,
)
from _text_utils import _DEFAULT_OCR_LANG, _fix_rtl_text

_IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tiff", ".tif", ".gif")


def read_pdf(
    path: str,
    pages: int = 0,
    ocr: bool = False,
    auto_ocr: bool = False,
    lang: str = _DEFAULT_OCR_LANG,
    engine_name: str = "auto",
    aggressive_clean: bool = False,
):
    """Read PDF with optional OCR fallback for scanned documents (F2).

    auto_ocr=False (default) - use --ocr flag explicitly for scanned PDFs.
    Set auto_ocr=True to enable automatic detection (may cause delays on large files).
    lang: OCR language(s) for Tesseract (default multi-language).
    engine_name: OCR backend ('auto' | 'tesseract'; both resolve to Tesseract 5.x).
    aggressive_clean: when True, run legacy heuristic transformations
        (`_fix_rtl_text` + `_dedup_word`) intended for PDFs that ship Hebrew
        glyphs in visual order. Modern pdfplumber returns logical-order
        Hebrew, where these transforms *corrupt* native double-letter words
        (e.g. "שלום", "ייחודי"). Default is False — preserve raw text.

    Size caps (env-configurable):
      - SENTINEL_PDF_MAX_BYTES   (default 50MB)
      - SENTINEL_PDF_MAX_PAGES   (default 1000)
    """
    err = _check_pdf_size(path)
    if err:
        return err

    result = _extract_pymupdf(path, pages, lang, aggressive_clean)
    if result is not None:
        return wrap_untrusted_content(result, source_name=f"PDF {path}")

    result = _extract_pdfplumber(
        path, pages, lang, aggressive_clean, ocr, auto_ocr, engine_name
    )
    if result is not None:
        return wrap_untrusted_content(result, source_name=f"PDF {path}")

    return wrap_untrusted_content(
        _extract_pypdf2(path, pages, aggressive_clean, ocr, auto_ocr),
        source_name=f"PDF {path}",
    )


def read_docx(path: str, as_markdown: bool = False):
    """Extract DOCX. When as_markdown=True, preserve headings, lists, and tables."""
    try:
        import docx
    except ImportError:
        return "❌ python-docx not installed."
    try:
        doc = docx.Document(path)
        if not as_markdown:
            return "\n".join(p.text for p in doc.paragraphs)
        parts = []
        for p in doc.paragraphs:
            style = p.style.name if p.style else "Normal"
            text = p.text.strip()
            if not text:
                continue
            if style.startswith("Heading"):
                level = style.replace("Heading", "").strip() or "1"
                parts.append(f"{'#' * int(level)} {text}")
            elif "List" in style or "Bullet" in style:
                parts.append(f"- {text}")
            else:
                parts.append(text)
        for table in doc.tables:
            parts.append("")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
            parts.append("")
        return wrap_untrusted_content("\n".join(parts), source_name=f"DOCX {path}")
    except Exception as e:
        return f"❌ DOCX read error: {e}"


def read_csv(path: str, query: str = ""):
    try:
        import pandas as pd
    except ImportError:
        return "❌ pandas not installed."
    try:
        df = pd.read_csv(path)
        lines = [f"Shape: {df.shape}", f"Columns: {list(df.columns)}"]
        if query:
            query_lower = query.lower()
            if "sample" in query_lower or "head" in query_lower:
                lines.append(df.head(5).to_string(index=False))
            elif "summary" in query_lower or "describe" in query_lower:
                lines.append(df.describe().to_string())
            elif "columns" in query_lower or "fields" in query_lower:
                lines.append(str(list(df.columns)))
            elif "count" in query_lower or "rows" in query_lower:
                lines.append(f"Total rows: {len(df)}")
            else:
                lines.append(df.head(10).to_string(index=False))
        else:
            lines.append(df.head(10).to_string(index=False))
        return wrap_untrusted_content("\n".join(lines), source_name=f"CSV {path}")
    except Exception as e:
        return f"❌ CSV read error: {e}"


def read_xlsx(path: str, query: str = ""):
    try:
        import pandas as pd
    except ImportError:
        return "❌ pandas not installed."
    try:
        xls = pd.ExcelFile(path)
        out = [f"Sheets: {xls.sheet_names}"]
        target = None
        if query:
            q_lower = query.lower()
            for name in xls.sheet_names:
                if name.lower() in q_lower:
                    target = name
                    break
        sheet = target or xls.sheet_names[0]
        df = pd.read_excel(path, sheet_name=sheet)
        out.append(f"Active sheet: {sheet}")
        out.append(f"Shape: {df.shape}")
        out.append(f"Columns: {list(df.columns)}")
        if query:
            q_lower = query.lower()
            if "sample" in q_lower or "head" in q_lower:
                out.append(df.head(5).to_string(index=False))
            elif "summary" in q_lower or "describe" in q_lower:
                out.append(df.describe().to_string())
            elif "columns" in q_lower or "fields" in q_lower:
                out.append(str(list(df.columns)))
            elif "count" in q_lower or "rows" in q_lower:
                out.append(f"Total rows: {len(df)}")
            else:
                out.append(df.head(10).to_string(index=False))
        else:
            out.append(df.head(10).to_string(index=False))
        return wrap_untrusted_content("\n".join(out), source_name=f"XLSX {path}")
    except Exception as e:
        return f"❌ XLSX read error: {e}"


def read_json(path: str):
    with open(path, encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, list):
        text = f"List with {len(data)} items.\nFirst item: {data[0] if data else 'empty'}"
    else:
        text = f"Keys: {list(data.keys())}"
    return wrap_untrusted_content(text, source_name=f"JSON {path}")


def read_txt(path: str):
    with open(path, encoding="utf-8", errors="replace") as f:
        return wrap_untrusted_content(f.read(), source_name=f"TXT {path}")
